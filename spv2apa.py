import pandas as pd
import docx
import re
from numpy import nan

input = "input.xlsx"
output = "output.docx"

# Check what we are dealing with, clean garbage data
df = pd.read_excel(input, nrows=0).columns[0]
df2 = pd.read_excel(input, nrows=1).iloc[0,0]

output_type = ""
analysis_type = ""
if df == "Coefficientsa":
    output_type = "Hierarchical Regression"
    df = pd.read_excel(input, header=2)
elif df == "Parameter Estimates":
    if df2 == "Parameter":
        output_type ="GeneralizedLM"
        analysis_type = "Univariate"
        df = pd.read_excel(input, header=1)
    elif "Dependent Variable:" in df2:
        output_type = "GLM"
        analysis_type = "Univariate"
        prefetch_dv = re.sub("Dependent Variable:   ", "", df2)
        df = pd.read_excel(input, header=2)
    else:
        output_type = "GLM"
        analysis_type = "Multivariate"
        df = pd.read_excel(input, header=1)
elif df == "Correlations":
    output_type = "Correlations"
    df = pd.read_excel(input, header=1)
    df = df.replace({"\*": "", ",": "0."}, regex=True)
    # Spearman correlations outputs have different layouts than Pearons, this fixes that
    if df.iloc[0,0] == "Spearman's rho":
        del df['Unnamed: 0']
        df = df.rename(columns={"Unnamed: 1": "Unnamed: 0", "Unnamed: 2": "Unnamed: 1"})

df = df.replace({".": nan, "0a": nan})

doc = docx.Document("template.docx")

# Significance is used to define the reference values to be used, check thresholds
thresholds = {1: [0.01, 0.05, 0.1],
              2: [0.001, 0.01, 0.05]}
significance = 2
p_notes = "Notes: " + "*** p < " + str(thresholds[significance][0]) + "; " + "** p < " + str(thresholds[significance][1]) + "; " + "* p < " + str(thresholds[significance][2])

def dataframe_to_docx(dataframe):
    # Source: https://stackoverflow.com/questions/40596518/writing-a-python-pandas-dataframe-to-word-document/40597684
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(dataframe.shape[0] + 1, dataframe.shape[1])
    t.style = "Table Grid"
    # add the header rows.
    for j in range(dataframe.shape[-1]):
        t.cell(0, j).text = dataframe.columns[j]

    # add the rest of the data frame
    for i in range(dataframe.shape[0]):
        for j in range(dataframe.shape[-1]):
            cell_value = str(dataframe.values[i, j])
            if cell_value == "nan":
                cell_value = ""
            t.cell(i + 1, j).text = cell_value

    doc.add_paragraph(p_notes)
    doc.save(output)

# Takes a p-value and returns asterisks
def sig_to_asterisks(p):
    cutoff = thresholds[significance]

    if p <= cutoff[0]:
        asterisks = "***"
    elif p <= cutoff[1]:
        asterisks = "**"
    elif p <= cutoff[2]:
        asterisks = "*"
    else:
        asterisks = ""
    return asterisks


# Clean up the dataframe
if output_type == "Hierarchical Regression":
    df = df.rename(columns={"Unnamed: 0": "Model", "Unnamed: 1": "Variable", "Unnamed: 5": "t", "Unnamed: 6": "Sig."})
    df = df.drop(df.tail(1).index)
elif output_type == "GLM":
    df = df.rename(columns={"Dependent Variable": "Model", "Parameter": "Variable", "95% Confidence Interval": "95% Confidence Interval LB", "Unnamed: 7": "95% Confidence Interval UB"})
    #df = df.drop(df.tail(1).index) # WHY WAS THIS HERE?
    if df.tail(1).iloc[0,0] == "a Computed using alpha = ,05":
        df = df.drop(df.tail(1).index)
    if df.tail(1).iloc[0,0] == "b Computed using alpha = ,05":
        df = df.drop(df.tail(1).index)
    if df.tail(1).iloc[0,0] == "a This parameter is set to zero because it is redundant.":
        df = df.drop(df.tail(1).index)
    df = df.drop(df.head(1).index).reset_index(drop=True)
    if analysis_type == "Univariate":
        df["Model"] = prefetch_dv
elif output_type == "GeneralizedLM":
    df = df.rename(columns={"Parameter": "Variable", "95% Wald Confidence Interval": "95% Wald Confidence Interval LB", "Unnamed: 4": "95% Wald Confidence Interval UB", "Unnamed: 6": "df", "Unnamed: 7": "Sig.", "95% Wald Confidence Interval for Exp(B)": "95% Wald Confidence Interval for Exp(B) LB", "Unnamed: 10": "95% Wald Confidence Interval for Exp(B) UB"})
    df = df.drop(0, axis=0).reset_index(drop=True)
    df["Model"] = re.sub("Dependent Variable: ", "", df.iloc[-4,0])
    df = df.drop(df.tail(5).index)
elif output_type == "Correlations":
    df = df.rename(columns={"Unnamed: 0": "Variable", "Unnamed: 1": "Parameter"})
    df = df.drop(df.tail(2).index)

df_final = pd.DataFrame()

if output_type != "Correlations":
    model_list = []
    model = ""
    if analysis_type == "Multivariate" or output_type == "Hierarchical Regression":
        # Replace NAN in model column with model number, get number of models
        for i in range(0, len(df)):
            cell = df.iloc[i, 0]
            if pd.isnull(cell) == False:
                model = cell
                model_list.append(model)
            else:
                df.iloc[i, 0] = model
    # For Univariates we only have a single model so we just need to inject that one to the model list
    elif analysis_type == "Univariate":
        model = df.loc[0, "Model"]
        model_list.append(model)

    # Get list of variables from full model, inject those into the cleaned df
    df_subset = df[df["Model"] == model_list[-1]]
    df_final["Variable"] = df_subset["Variable"].reset_index(drop=True)

    # For each model, get parameters, format them, and add as new columns to the final df
    for i in model_list:
        cells_list = []
        df_subset = df[df["Model"] == i].reset_index(drop=True)

        for j in range(0, len(df_subset)):
            b = df_subset.loc[j, "B"]
            se = df_subset.loc[j, "Std. Error"]
            sig = df_subset.loc[j, "Sig."]
            # Check if this is a blank row such as a reference category, otherwise add the proper contents
            if pd.isnull(b) == True:
                cell_value = nan
            else:
                cell_value = "{0:.3f}".format(b) + sig_to_asterisks(sig) + "\n" + "(" + "{0:.3f}".format(se) + ")"
            cells_list.append(cell_value)

        df_final[str(i)] = pd.Series(cells_list)

    # Clear reference categories
    if output_type == "GLM" or output_type == "GeneralizedLM":
        df_final = df_final.dropna()

elif output_type == "Correlations":

    # For each row and each column, check if we're in the correlation row, and if yes, grab the p-value from lower row
    for i in range(0, len(df)):
        for j in range(2, len(df.columns)):
            cell = df.iloc[i, j]
            if df.iloc[i,1] == "Pearson Correlation" or df.iloc[i,1] == "Correlation Coefficient":
                sig = df.iloc[i+1, j]
                if pd.isnull(sig) == True:
                    pass
                else:
                    df.iloc[i,j] = str(cell) + sig_to_asterisks(sig)

    # Clean up unnecessary rows and columns
    df = df.dropna()
    df = df.drop(columns="Parameter")

    # Remove redundant lower half of the matrix
    for j in range(1, len(df.columns)):
        mirror = 0
        for i in range(0, len(df)):
            if mirror == 1:
                df.iloc[i,j] = ""
            if df.iloc[i,j] == 1:
                mirror = 1
    df_final = df

else:
    print("NO RECOGNIZED OUTPUT TYPES!")

dataframe_to_docx(df_final)
